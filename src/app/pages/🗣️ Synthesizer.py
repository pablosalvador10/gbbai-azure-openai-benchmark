import streamlit as st
import dotenv
import os
import requests
from utils.ml_logging import get_logger
from src.extractors.blob_data_extractor import AzureBlobDataExtractor
from src.aoai.azure_openai import AzureOpenAIManager
from src.ocr.document_intelligence import AzureDocumentIntelligenceManager
from docx import Document
from src.utilsfunc import save_uploaded_file

# Load environment variables
dotenv.load_dotenv(".env")

# Set up logger
logger = get_logger()

def markdown_to_docx(markdown_text):
    doc = Document()
    lines = markdown_text.split('\n')
    
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            paragraph = doc.add_paragraph(style='ListBullet')
            process_bold_text(line[2:], paragraph)
        else:
            paragraph = doc.add_paragraph()
            process_bold_text(line, paragraph)
    
    doc.save('Generated_Guide.docx')
    return doc

def process_bold_text(text, paragraph):
    while '**' in text:
        start_bold = text.find('**')
        end_bold = text.find('**', start_bold + 2)
        if start_bold != -1 and end_bold != -1:
            paragraph.add_run(text[:start_bold])
            paragraph.add_run(text[start_bold+2:end_bold]).bold = True
            text = text[end_bold+2:]
        else:
            break
    paragraph.add_run(text)

# Initialize managers in session state if not already present
for manager_name, manager in [
    ("document_intelligence_manager", AzureDocumentIntelligenceManager()),
    ("blob_data_extractor_manager", AzureBlobDataExtractor(container_name="ocrtest2")),
    ("azure_openai_manager", AzureOpenAIManager())
]:
    if manager_name not in st.session_state:
        st.session_state[manager_name] = manager

# Sidebar for inputs
with st.sidebar:
    st.title("🤖 AI RequestGPT")
    uploaded_files = st.file_uploader("Upload documents", type=["png", "jpg", "jpeg", "pdf"], accept_multiple_files=True)
    user_inputs = st.text_input("Enter your instructions")
    topic = st.text_input("Enter the topic")
    submit_to_ai = st.button("Submit to AI")

# Initialize session state variables if they don't exist
if 'conversation_history' not in st.session_state:
    st.session_state.conversation_history = []
if 'ai_response' not in st.session_state:
    st.session_state.ai_response = ""
if 'chat_history' not in st.session_state:  # Initialize chat_history
    st.session_state.chat_history = []

# Display only the latest AI response in an expandable result box
if "ai_response" in st.session_state and st.session_state["ai_response"]:
    with st.expander("AI Response", expanded=True):
        st.markdown(st.session_state["ai_response"], unsafe_allow_html=True)

# Function to generate AI response
def generate_ai_response(user_query):
    # Placeholder for the actual call to generate AI response
    # Replace this with the actual call to your AI model
    ai_response = st.session_state.azure_openai_manager.generate_chat_response(
        conversation_history=st.session_state.conversation_history,
        system_message_content='''You are tasked with creating detailed, user-friendly documentation based on multiple documents and complex topics. 
        The goal is to distill this information into an easy-to-follow "How-To" guide. This documentation should 
        be structured with clear headings, subheadings, and step-by-step instructions that guide the user through 
        the necessary processes or concepts. Each section should be well-organized and written in simple language 
        to ensure that the content is accessible and understandable to users with varying levels of expertise. 
        The documentation should cover the setup, configuration, and usage of tools or techniques, 
        including practical examples and troubleshooting tips to address common issues or challenges that users 
        might encounter.''',
        query=user_query,
        max_tokens=3000
    )
    return ai_response

# Process uploaded files and user inputs after submission
if submit_to_ai and uploaded_files and user_inputs and topic:
    
    markdown_content = ""
    with st.spinner("Processing uploaded files..."):
        for uploaded_file in uploaded_files:
            file_path = save_uploaded_file(uploaded_file)
            try:
                blob_url = st.session_state.blob_data_extractor_manager.upload_file_to_blob(file_path, uploaded_file.name)
                result_ocr = st.session_state.document_intelligence_manager.analyze_document(
                    document_input=blob_url,
                    model_type="prebuilt-layout",
                    output_format="markdown",
                    features=["OCR_HIGH_RESOLUTION"]
                )
                markdown_content += result_ocr.content + "\n\n"
                st.success(f"Document '{uploaded_file.name}' has been successfully processed.")
            except Exception as e:
                logger.error(f"Error processing file {uploaded_file.name}: {e}")
                st.error(f"Error processing file {uploaded_file.name}. Please check the logs for more details.")

    with st.spinner("Processing your request..."):
        max_tokens = 3000
        query = f'''Given the content extracted from various documents using Optical Character Recognition (OCR) technology and provided in markdown format, your task is to create a high-quality, detailed "How-To" guide. The guide should distill complex topics into accessible, step-by-step instructions tailored for users seeking to understand or implement specific processes or concepts.
            userinputs: {user_inputs}
            context: {markdown_content}

            Essential Steps for Crafting the Guide:

            1. **Content Synthesis**: Begin by synthesizing the OCR-extracted content. Identify crucial themes, technical concepts, and actionable instructions relevant to Copilot X and productivity enhancement. This synthesis forms the foundation of your guide's structure and content focus.

            2. **Target Audience Clarification**: Clearly define the guide's target audience. Understanding the audience's technical background, familiarity with Copilot X, and productivity goals is essential for customizing the guide's complexity and instructional style.

            3. **Structured Outline Development**: Construct a structured outline to organize the guide into coherent sections and subsections. Each section should concentrate on distinct aspects of using Copilot X for productivity, ensuring a logical progression from introductory concepts to advanced applications.

            4. **Guide Composition**:
                a. **Introduction**: Craft an introduction that outlines the guide's objectives, the significance of Copilot X for productivity, and what the readers will gain.
                b. **Detailed Instructions**: Following the outline, elaborate on each section with clear, technical instructions. Incorporate step-by-step processes, code snippets, examples, and best practices specific to Copilot X.
                c. **Conclusion**: Summarize the key takeaways, suggest further reading or resources, and encourage steps for practical application.

            5. **Comprehensive Review and Enhancement**: Thoroughly review the guide to ensure technical accuracy, clarity, and completeness. Revise any sections as necessary, and consider peer or expert feedback for additional insights.

            6. **Final Formatting and Release**: Apply professional formatting to enhance readability and visual appeal. Use diagrams, screenshots, or videos where applicable. Publish the guide in a format accessible to your target audience, ensuring it's ready for distribution and application.

            Additional Guidelines:

            - Begin with a clear agenda and systematically develop content within designated sections.
            - Employ straightforward language while explaining technical details, using examples to demystify complex concepts.
            - Dedicate ample time to crafting high-quality content, prioritizing accuracy and user engagement.
            - Base the guide explicitly on the OCR content and the nuanced requirements of the user's query regarding {topic}.
            - minimun length of the document should be {max_tokens} tokens'''
        
        st.session_state.chat_history.append({"role": "user", "content": query})
        try:
            ai_response = generate_ai_response(query)
            st.session_state['ai_response'] = ai_response
            st.session_state.chat_history.append({"role": "ai", "content": ai_response})
        except Exception as e:
            logger.error(f"Error communicating with Azure OpenAI: {e}")
            st.error("Error communicating with Azure OpenAI. Please check the logs for more details.")

# Initialize chat history
if "messages" not in st.session_state:
    st.session_state.messages = []

# Display chat messages from history on app rerun
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Corrected and improved version of the code snippet

if 'messages' not in st.session_state:
    st.session_state['messages'] = []

if prompt := st.text_input("How can I assist you today?"):
    # Add user message to chat history
    st.session_state.messages.append({"role": "user", "content": prompt})

    # Prepare conversation history for the AI model, excluding system messages
    conversation_history = [message for message in st.session_state.messages if message["role"] != "system"]

    # Generate AI response using the conversation history
    ai_response = st.session_state.azure_openai_manager.generate_chat_response(
        conversation_history=conversation_history,
        system_message_content='''You are tasked with creating detailed, user-friendly documentation based on multiple documents and complex topics. 
        The goal is to distill this information into an easy-to-follow "How-To" guide. This documentation should 
        be structured with clear headings, subheadings, and step-by-step instructions that guide the user through 
        the necessary processes or concepts. Each section should be well-organized and written in simple language 
        to ensure that the content is accessible and understandable to users with varying levels of expertise. 
        The documentation should cover the setup, configuration, and usage of tools or techniques, 
        including practical examples and troubleshooting tips to address common issues or challenges that users 
        might encounter.''',
        query=prompt,
        max_tokens=3000
    )
    # Update the AI response in session state
    st.session_state['ai_response'] = ai_response

    # Add assistant response to chat history
    st.session_state.messages.append({"role": "assistant", "content": ai_response})

    # Display the conversation history
    for message in st.session_state.messages:
        role = message["role"]
        content = message["content"]
        with st.container():
            if role == "user":
                st.markdown(f"**You:** {content}")
            else:
                st.markdown(f"**Assistant:** {content}")

# Download button for the generated guide
if st.button('Download Guide as DOCX'):
    ai_response = st.session_state.get("ai_response", "")
    if ai_response:
        with st.spinner("Generating DOCX file..."):
            try:
                doc = markdown_to_docx(ai_response)
                with open('Generated_Guide.docx', 'rb') as file:
                    st.download_button('Download Guide', file, 'Generated_Guide.docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                st.success("Guide generated and ready for download.")
            except Exception as e:
                logger.error(f"Error generating DOCX file: {e}")
                st.error("Error generating DOCX file. Please check the logs for more details.")
    else:
        st.error("No AI response available to generate the DOCX file.")